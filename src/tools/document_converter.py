"""
Document Conversion Utility
---------------------------

This module provides robust document-to-Markdown conversion using a multi-tiered
fallback strategy for maximum reliability and quality.

WHY THIS MODULE?
- Universal Support: Handles PDF, DOCX, MD, and TXT files
- Best Quality: Uses Docling for superior PDF parsing (multi-column, tables, complex layouts)
- Reliability: Falls back gracefully if advanced libraries are unavailable
- Zero-Config: Works out of the box with intelligent library detection

CONVERSION STRATEGY:
1. Docling (Tier 1): Advanced PDF/DOCX parsing - handles complex layouts, tables, headers
2. PyPDF2/python-docx (Tier 2): Basic parsing - reliable fallback for simple documents
3. Plain text: Direct read for .txt and .md files

DESIGN PRINCIPLES:
- Graceful Degradation: Try best method first, fall back if needed
- Fail-Safe: Always return content, even if quality is reduced
- Transparent Logging: User knows which method succeeded and why
- Type Safety: Returns clean, validated Markdown strings
"""

from pathlib import Path

# Handle imports for both package usage and direct script execution
try:
    from src.core.logger import get_logger
except ImportError:
    # Fallback for when running this file directly
    import sys
    sys.path.insert(0, str(Path(__file__).parent.parent.parent))
    from src.core.logger import get_logger

logger = get_logger(__name__)


# ==============================================================================
# Main Conversion Function
# ==============================================================================

def convert_document_to_markdown(file_path: str) -> str:
    """
    Convert any supported document format to clean Markdown text.
    
    This is the main entry point for document conversion. It intelligently selects
    the best conversion method based on file format and available libraries.
    
    Supported Formats:
    - PDF (.pdf): Uses Docling → PyPDF2 fallback chain
    - DOCX (.docx): Uses Docling → python-docx fallback chain
    - Markdown (.md): Direct read (no conversion needed)
    - Plain Text (.txt): Direct read (no conversion needed)
    
    Args:
        file_path: Path to the document file to convert.
                  Can be absolute or relative path.
    
    Returns:
        Clean Markdown text extracted from the document.
        Always returns a string, never None.
    
    Raises:
        FileNotFoundError: If the file does not exist at the specified path.
        ValueError: If the file format is not supported.
        IOError: If all conversion methods fail (rare, usually means file is corrupt).
    
    Example:
        >>> markdown = convert_document_to_markdown("resume.pdf")
        >>> print(f"Extracted {len(markdown)} characters")
        Extracted 5234 characters
    """
    # Convert to Path object for easier manipulation
    path = Path(file_path)
    
    # Validate that the file exists
    if not path.is_file():
        error_msg = f"File not found: {file_path}"
        logger.error(error_msg)
        raise FileNotFoundError(error_msg)
    
    # Get the file extension (lowercase for case-insensitive comparison)
    file_extension = path.suffix.lower()
    
    # Handle plain text and Markdown files directly (no conversion needed)
    if file_extension in [".md", ".txt"]:
        return _read_text_file(path)
    
    # Handle PDF files with fallback chain
    if file_extension == ".pdf":
        return _convert_pdf_with_fallback(path)
    
    # Handle DOCX files with fallback chain
    if file_extension == ".docx":
        return _convert_docx_with_fallback(path)
    
    # Unsupported file format
    supported_formats = [".pdf", ".docx", ".md", ".txt"]
    error_msg = (
        f"Unsupported file format: {file_extension}. "
        f"Supported formats: {', '.join(supported_formats)}"
    )
    logger.error(error_msg)
    raise ValueError(error_msg)


# ==============================================================================
# Plain Text Reading
# ==============================================================================

def _read_text_file(path: Path) -> str:
    """
    Read a plain text or Markdown file directly.
    
    Args:
        path: Path object pointing to the text file.
    
    Returns:
        File contents as a string.
    
    Raises:
        IOError: If file reading fails.
    """
    logger.info(f"Reading text file directly: {path.name}")
    
    try:
        content = path.read_text(encoding="utf-8")
        logger.debug(f"Successfully read {len(content)} characters from {path.name}")
        return content
    
    except Exception as e:
        error_msg = f"Failed to read text file {path}: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise IOError(error_msg) from e


# ==============================================================================
# PDF Conversion with Fallback Chain
# ==============================================================================

def _convert_pdf_with_fallback(path: Path) -> str:
    """
    Convert PDF to Markdown using Docling (primary and only method).
    
    This function uses Docling for PDF conversion. Docling provides superior
    quality for complex layouts, tables, and multi-column documents.
    
    Args:
        path: Path object pointing to the PDF file.
    
    Returns:
        Markdown content extracted from the PDF.
    
    Raises:
        IOError: If Docling conversion fails.
    """
    # Use Docling (best quality - handles complex layouts, tables, multi-column)
    try:
        content = _convert_pdf_with_docling(path)
        logger.info(f"Successfully converted PDF with Docling: {path.name}")
        return content
    except ImportError as e:
        error_msg = (
            f"Docling is not installed. Install with: uv pip install docling. "
            f"Original error: {str(e)}"
        )
        logger.error(error_msg)
        raise IOError(error_msg) from e
    except Exception as e:
        error_msg = str(e)
        # Handle Unicode encoding issues in error messages for Windows console
        try:
            error_msg = error_msg.encode('utf-8', errors='replace').decode('utf-8')
        except Exception:
            error_msg = repr(e)
        
        # Check for Windows symlink permission error
        if "WinError 1314" in error_msg or "required privilege" in error_msg.lower():
            full_error = (
                f"Docling PDF conversion failed due to Windows permission issue. "
                f"Docling requires symlink support. Solutions: "
                f"1) Enable Windows Developer Mode, or "
                f"2) Run Python as administrator, or "
                f"3) Set environment variable: HF_HUB_DISABLE_SYMLINKS_WARNING=1. "
                f"File: {path.name}"
            )
        else:
            full_error = (
                f"Docling PDF conversion failed for {path.name}. "
                f"Error: {error_msg}. "
                f"Please ensure Docling is properly installed: uv pip install docling"
            )
        logger.error(full_error, exc_info=True)
        raise IOError(full_error) from e


def _convert_pdf_with_docling(path: Path) -> str:
    """
    Convert PDF to Markdown using Docling (advanced method).
    
    Docling excels at:
    - Multi-column layouts (academic papers, newsletters)
    - Tables and structured data
    - Headers and footers (automatically filtered)
    - Complex formatting with semantic structure preservation
    - Scanned documents (OCR support)
    
    Args:
        path: Path object pointing to the PDF file.
    
    Returns:
        Clean Markdown content extracted by Docling.
    
    Raises:
        ImportError: If Docling is not installed.
        Exception: If Docling processing fails.
    """
    try:
        from docling.document_converter import DocumentConverter
        
        logger.info(f"Converting PDF with Docling: {path.name}")
        
        # Initialize the Docling converter
        converter = DocumentConverter()
        
        # Convert the document
        # Docling returns a ConversionResult object
        result = converter.convert(str(path))
        
        # Extract the Markdown content
        # The document object has an export_to_markdown() method
        markdown_content = result.document.export_to_markdown()
        
        logger.debug(f"Docling extracted {len(markdown_content)} characters")
        return markdown_content
    
    except ImportError as e:
        logger.debug("Docling library not installed")
        raise ImportError("Docling is not installed. Install with: uv pip install docling") from e
    
    except Exception as e:
        error_msg = str(e)
        # Handle Unicode encoding issues in error messages
        try:
            error_msg = error_msg.encode('utf-8', errors='replace').decode('utf-8')
        except Exception:
            error_msg = repr(e)
        logger.debug(f"Docling conversion error: {error_msg}")
        raise Exception(f"Docling processing failed: {error_msg}") from e


def _convert_pdf_with_pypdf2(path: Path) -> str:
    """
    Convert PDF to text using PyPDF2 (basic method).
    
    Limitations:
    - No table extraction
    - No multi-column support
    - Minimal structure preservation
    - May miss headers/footers
    - No OCR for scanned documents
    
    Args:
        path: Path object pointing to the PDF file.
    
    Returns:
        Plain text content with basic Markdown formatting.
    
    Raises:
        ImportError: If PyPDF2 is not installed.
        Exception: If PDF extraction fails.
    """
    try:
        import PyPDF2
        
        logger.info(f"Converting PDF with PyPDF2: {path.name}")
        
        text_pages = []
        
        with open(path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                try:
                    text = page.extract_text()
                    if text.strip():  # Only add non-empty pages
                        text_pages.append(f"# Page {page_num}\n\n{text}\n")
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")
        
        if not text_pages:
            raise Exception("No text content extracted from PDF (file may be scanned or empty)")
        
        markdown_content = "\n".join(text_pages)
        logger.debug(f"PyPDF2 extracted {len(markdown_content)} characters")
        return markdown_content
    
    except ImportError as e:
        raise ImportError("PyPDF2 is not installed. Install with: uv pip install PyPDF2") from e
    
    except Exception as e:
        raise Exception(f"PyPDF2 extraction failed: {str(e)}") from e


# ==============================================================================
# DOCX Conversion with Fallback Chain
# ==============================================================================

def _convert_docx_with_fallback(path: Path) -> str:
    """
    Convert DOCX to Markdown using fallback chain: Docling → python-docx.
    
    Args:
        path: Path object pointing to the DOCX file.
    
    Returns:
        Markdown content extracted from the DOCX.
    
    Raises:
        IOError: If all conversion methods fail.
    """
    # Try Docling first (best quality)
    try:
        content = _convert_docx_with_docling(path)
        logger.info(f"✅ Successfully converted DOCX with Docling: {path.name}")
        return content
    except ImportError:
        logger.debug("Docling not available, trying python-docx fallback")
    except Exception as e:
        logger.warning(f"Docling conversion failed: {e}, trying python-docx fallback")
    
    # Fall back to python-docx (basic but reliable)
    try:
        content = _convert_docx_with_python_docx(path)
        logger.info(f"✅ Successfully converted DOCX with python-docx (fallback): {path.name}")
        return content
    except Exception as e:
        error_msg = f"All DOCX conversion methods failed for {path}: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise IOError(error_msg) from e


def _convert_docx_with_docling(path: Path) -> str:
    """
    Convert DOCX to Markdown using Docling (advanced method).
    
    Args:
        path: Path object pointing to the DOCX file.
    
    Returns:
        Clean Markdown content extracted by Docling.
    
    Raises:
        ImportError: If Docling is not installed.
        Exception: If Docling processing fails.
    """
    try:
        from docling.document_converter import DocumentConverter
        
        logger.info(f"Converting DOCX with Docling: {path.name}")
        
        converter = DocumentConverter()
        result = converter.convert(str(path))
        markdown_content = result.document.export_to_markdown()
        
        logger.debug(f"Docling extracted {len(markdown_content)} characters")
        return markdown_content
    
    except ImportError as e:
        raise ImportError("Docling is not installed. Install with: uv pip install docling") from e
    
    except Exception as e:
        raise Exception(f"Docling processing failed: {str(e)}") from e


def _convert_docx_with_python_docx(path: Path) -> str:
    """
    Convert DOCX to Markdown using python-docx (basic method).
    
    Args:
        path: Path object pointing to the DOCX file.
    
    Returns:
        Markdown content with basic formatting.
    
    Raises:
        ImportError: If python-docx is not installed.
        Exception: If DOCX extraction fails.
    """
    try:
        from docx import Document
        
        logger.info(f"Converting DOCX with python-docx: {path.name}")
        
        doc = Document(str(path))
        text_parts = []
        
        # Extract paragraphs with basic heading detection
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect headings by style name
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level_num = int(level)
                        text_parts.append(f"{'#' * level_num} {text}\n")
                    except ValueError:
                        text_parts.append(f"{text}\n")
                else:
                    text_parts.append(f"{text}\n")
        
        # Extract tables
        for table in doc.tables:
            text_parts.append("\n**Table:**\n")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                text_parts.append(f"{row_text}\n")
            text_parts.append("\n")
        
        if not text_parts:
            raise Exception("No text content extracted from DOCX (file may be empty)")
        
        markdown_content = "\n".join(text_parts)
        logger.debug(f"python-docx extracted {len(markdown_content)} characters")
        return markdown_content
    
    except ImportError as e:
        raise ImportError("python-docx is not installed. Install with: uv pip install python-docx") from e
    
    except Exception as e:
        raise Exception(f"python-docx extraction failed: {str(e)}") from e


# ==============================================================================
# Utility Functions
# ==============================================================================

def get_supported_formats() -> list[str]:
    """
    Get list of supported document formats.
    
    Returns:
        List of supported file extensions (e.g., ['.pdf', '.docx', '.md', '.txt']).
    """
    return [".pdf", ".docx", ".md", ".txt"]


def is_format_supported(file_path: str) -> bool:
    """
    Check if a file format is supported.
    
    Args:
        file_path: Path to the file to check.
    
    Returns:
        True if the file format is supported, False otherwise.
    
    Example:
        >>> is_format_supported("resume.pdf")
        True
        >>> is_format_supported("resume.jpg")
        False
    """
    path = Path(file_path)
    return path.suffix.lower() in get_supported_formats()


def get_available_converters() -> dict[str, bool]:
    """
    Check which conversion libraries are currently available.
    
    Returns:
        Dictionary mapping converter names to availability status.
    
    Example:
        >>> converters = get_available_converters()
        >>> print(converters)
        {'docling': True, 'pypdf2': True, 'python-docx': True}
    """
    converters = {}
    
    # Check for Docling
    try:
        import docling
        converters["docling"] = True
    except ImportError:
        converters["docling"] = False
    
    # Check for PyPDF2
    try:
        import PyPDF2
        converters["pypdf2"] = True
    except ImportError:
        converters["pypdf2"] = False
    
    # Check for python-docx
    try:
        import docx
        converters["python-docx"] = True
    except ImportError:
        converters["python-docx"] = False
    
    return converters


# ==============================================================================
# Testing Block
# ==============================================================================

if __name__ == "__main__":
    """
    Test the document converter with sample files.
    Run this script directly to test document conversion.
    """
    print("=" * 70)
    print("Document Converter Test")
    print("=" * 70)

    # Check which converters are available
    converters = get_available_converters()
    print("\n--- Available Converters ---")
    for name, available in converters.items():
        status = "[OK] Available" if available else "[X] Not installed"
        print(f"{name}: {status}")

    # List supported formats
    print("\n--- Supported Formats ---")
    print(", ".join(get_supported_formats()))

    print("\n" + "=" * 70)
    print("To test with a file, call:")
    print("  convert_document_to_markdown('path/to/your/file.pdf')")
    print("=" * 70)

