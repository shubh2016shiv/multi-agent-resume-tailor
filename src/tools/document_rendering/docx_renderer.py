"""Resume -> .docx (Microsoft Word): the format recruiters and many ATS expect.

Uses python-docx (pure Python, cross-platform -- no system toolchain, unlike the PDF
path), so a .docx is always produced wherever the project installs. Mirrors the markdown
and PDF layout: the same profile-based section order (section_policy) and the same date
format, so all output formats present the candidate identically.
"""

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument

from src.data_models.resume import Resume
from src.tools.document_rendering.section_policy import (
    ResumeSection,
    experience_date_range,
    group_skills,
    infer_profile,
    section_order,
)


def render_resume_docx(resume: Resume, output_path: Path) -> Path:
    """Write the resume as a .docx at output_path and return it (always succeeds).

    Section order follows the same career-stage profile as the PDF/Markdown; empty
    sections add nothing.
    """
    document = Document()
    _add_header(document, resume)
    section_adders = {
        ResumeSection.SUMMARY: _add_summary,
        ResumeSection.SKILLS: _add_skills,
        ResumeSection.EXPERIENCE: _add_experience,
        ResumeSection.EDUCATION: _add_education,
        ResumeSection.CERTIFICATIONS: _add_certifications,
        ResumeSection.LANGUAGES: _add_languages,
    }
    for section in section_order(infer_profile(resume)):
        section_adders[section](document, resume)
    document.save(str(output_path))
    return output_path


def _add_header(document: DocxDocument, resume: Resume) -> None:
    """The name as the title, then a single contact line."""
    document.add_heading(resume.full_name, level=0)
    contact = [resume.email]
    if resume.phone_number:
        contact.append(resume.phone_number)
    if resume.location:
        contact.append(resume.location)
    if resume.website_or_portfolio:
        contact.append(resume.website_or_portfolio)
    document.add_paragraph(" | ".join(contact))


def _add_summary(document: DocxDocument, resume: Resume) -> None:
    if not resume.professional_summary:
        return
    document.add_heading("Summary", level=1)
    document.add_paragraph(resume.professional_summary)


def _add_skills(document: DocxDocument, resume: Resume) -> None:
    if not resume.skills:
        return
    document.add_heading("Skills", level=1)
    for category, names in group_skills(resume.skills):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{category}: ").bold = True
        paragraph.add_run(", ".join(names))


def _add_experience(document: DocxDocument, resume: Resume) -> None:
    if not resume.work_experience:
        return
    document.add_heading("Experience", level=1)
    for role in resume.work_experience:
        title = document.add_paragraph()
        title.add_run(f"{role.job_title}, {role.company_name}").bold = True
        meta = " | ".join(part for part in (role.location, experience_date_range(role)) if part)
        if meta:
            document.add_paragraph().add_run(meta).italic = True
        for achievement in role.achievements:
            document.add_paragraph(achievement, style="List Bullet")


def _add_education(document: DocxDocument, resume: Resume) -> None:
    if not resume.education:
        return
    document.add_heading("Education", level=1)
    for entry in resume.education:
        heading = document.add_paragraph()
        heading.add_run(f"{entry.institution_name} ({entry.graduation_year})").bold = True
        degree = entry.degree
        if entry.field_of_study:
            degree += f" in {entry.field_of_study}"
        document.add_paragraph(degree)


def _add_certifications(document: DocxDocument, resume: Resume) -> None:
    if not resume.certifications:
        return
    document.add_heading("Certifications", level=1)
    for item in resume.certifications:
        document.add_paragraph(item, style="List Bullet")


def _add_languages(document: DocxDocument, resume: Resume) -> None:
    if not resume.languages:
        return
    document.add_heading("Languages", level=1)
    document.add_paragraph(", ".join(resume.languages))
